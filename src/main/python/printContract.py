#!python3
# printContract.py - prints a contract

import sys, pprint, docx, os
from docx import Document
from docx.shared import Pt

# initialize values given from cli
eventType = sys.argv[1] 
name = sys.argv[2] 
contact = sys.argv[3] 
venue = sys.argv[4] 
signUpDate = sys.argv[5] 
eventDate = sys.argv[6] 
packageInclusion = sys.argv[7] 
price = sys.argv[8] 
clientBudget = sys.argv[9]
notes = sys.argv[10]
directory = "C:\\Users\\" + os.getlogin() + "\\Documents"

class PrintContract(object):

    # initalize variables with parameters
    def __init__(self, eventType, name, contact, venue, signUpDate, eventDate, packageInclusion, price, clientBudget, notes):
        self.eventType = eventType
        self.name = name
        self.contact = contact
        self.venue = venue
        self.signUpDate = signUpDate
        self.eventDate = eventDate
        self.packageInclusion = packageInclusion
        self.price = price
        self.clientBudget = clientBudget
        self.notes = notes

        self.doc = Document()

        ### 
        os.chdir(directory)
        self.initializeElements()
        self.generateFile()


    # initialize program
    def initializeElements(self):
        self.title = "EventsByRhandell"
        self.paragraph1 = f'''
    This contract is made effective as of {self.signUpDate} by and between Eventsbyrhandell and {self.name}. This contract serves as an agreement of distribution of services and compensation in the planning of {self.eventType}, which will be held on {self.eventDate}. 
    '''

        self.paragraph2 = f'''The Event Planner is responsible for the planning and coordination of the Event. This shall include making appropriate reservations, creating a format for invitations, reserving necessary space for the Event, and handling on-site logistics. Additional duties may be required based on the nature of the Event and should be described below:
{self.notes}'''

        self.paragraph3 = f'''As compensation for the Event Planner’s services, the Client agrees to pay a total of [Amount] pesos to the Event Planner. A non-refundable deposit of 25% or [Amount] pesos will be made as an initial deposit and is due on the date that this contract goes into effect. 
The remaining balance of {self.price} pesos shall be due no later than five (5) business days prior to the date of the event. If the balance is not received within this period, an interest rate of 15% of the remaining balance shall be charged for every week the balance is past due.'''

        self.paragraph4 = '''Based on preliminary discussions, the Event Planner will compile a list of preferred vendors and review that list with the Client. This list may include, but is not limited to, caterers, rental agencies, musicians, DJs, photographers, Videographers, chefs, hotels, and conference centers. Services desired by the Client will be drawn from the approved vendor list when possible. Deviations from this list will be announced to the Client prior to the approval of any final agreement with the vendor. 
The Client is responsible for all payments made to the vendors chosen, and will be notified of all due dates at the time that an agreement is formed with these parties.'''

        self.paragraph5 = '''Should the date of the Event change, the Event Planner will make the best effort possible to accommodate the new date. The Client understands that last-minute changes can affect the quality of the final Event and that these changes are not necessarily the fault of the Event Planner. 
In the event of a cancellation of the Event, the Client should notify the Event Planner no later than thirty (30) days prior to the planned date. Should the Event be canceled after that deadline, the Event Planner may collect the full fee owed, including any applicable interest that may arise due to late payments.'''
        self.paragraph6 = '''The Event Planner will try to select an alternate venue for outdoor events where weather is a concern if possible. In the event of severe weather that may disrupt the Event or prevent it from being held, the Client can propose an alternate date and the Event Planner will make the best effort to accommodate that change as per the clause above. '''

        self.paragraph7 = f'''By signing below, both the Client and the Event Planner indicate that they have read, understand, and agree to all terms and conditions presented above.


Client:  {self.name}                  
Date:	

Event Planner:	Rhandell Villados	          
Date:	

        '''
        
        # dict initalization for headings and paragraphs
        self.headings = {
            1:"\n\nScope of Services",
            2:"\n\nCompensation",
            3:"\n\nVendors",
            4:"\n\nDate Changes and Cancellations",
            5:"\n\nWeather"
        }
        self.paragraphs = {
            1:self.paragraph2,
            2:self.paragraph3,
            3:self.paragraph4,
            4:self.paragraph5,
            5:self.paragraph6
        }


        # set the styles
        self.style = self.doc.styles['Normal']
        self.font = self.style.font
        self.font.name = "Times New Roman"
        self.font.size = Pt(12)

        self.doc.add_heading("EventsByRhandell", level=0)
        self.docPar = self.doc.add_paragraph(self.paragraph1)

        # the adding of elements
        for i in range(1, len(self.headings)):
            self.docPar.add_run(self.headings[i]).bold = True
            self.docPar = self.doc.add_paragraph(self.paragraphs[i])
        self.doc.add_paragraph(self.paragraph7)

    # generate the file 
    def generateFile(self):
        self.doc.save(self.name + "Contract.docx")
        print("Done!")

# create an instance
run = PrintContract(
    eventType, 
    name,
    contact, 
    venue, 
    signUpDate, 
    eventDate, 
    packageInclusion, 
    price, 
    clientBudget,
    notes)